#!/usr/bin/python
# -*- coding: utf-8 -*-
"""
linkconverter.py

This program is a template for python programs
All this stuff at the top of the script is just optional metadata;

"""
import docx
from docx.oxml.shared import OxmlElement, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.enum.dml import MSO_THEME_COLOR_INDEX


__author__ = "Denis J Jackman (denis_jackman@hotmail.com)"
__version__ = "$Revision: 0.00 $"
__date__ = "$Date: 2022/09/12 09:00:00 $"
__copyright__ = "Copyright (c) 2022 Denis J Jackman"
__license__ = "Python"

def add_hyperlink(document, paragraph, url, name):
    """
    Add a hyperlink to a paragraph.

    :param document: The Document being edited.
    :param paragraph: The Paragraph the hyperlink is being added to.
    :param url: The url to be added to the link.
    :param name: The text for the link to be displayed in the paragraph
    :return: None
    """

    part = document.part
    rId = part.relate_to(url, RT.HYPERLINK, is_external=True)

    init_hyper = OxmlElement('w:hyperlink')
    init_hyper.set(qn('r:id'), rId, )
    init_hyper.set(qn('w:history'), '1')

    new_run = OxmlElement('w:r')

    rPr = OxmlElement('w:rPr')

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')

    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = name
    init_hyper.append(new_run)

    r = paragraph.add_run()
    r._r.append(init_hyper)
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True



def main():
    """ This is the main routine for the program """
    print("Starting the sequence:")
    document = docx.Document()

    with open('y:/pyproject/scratch/example.txt', 'r', encoding='utf8') as file:
        for line in file:
            para = document.add_paragraph()
            para.style = 'List Bullet'
            item = line.split('/')
            textlink = item[-1].strip()
            add_hyperlink(document, para, line, textlink)
            para.add_run("\n")

    document.save('demo.docx')
    print("finishing up and closing down:")


if __name__ == '__main__':
    main()
