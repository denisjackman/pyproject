'''
    extract stuff from a pdf file
'''
import os
import sys
import docx
from tqdm import tqdm
# pylint: disable=C0413
sys.path.append(os.path.realpath('..\\..'))
from jackmanimation.utilities.fileutility import walk_through  # noqa: flake8 E402, E501
from jackmanimation.utilities.fileutility import extract_file_extension  # noqa: flake8 E402, E501
TARGETDIRECTORY = r"G:\PPB\Projects\migration status"
TARGETSTRING = "Exchange"


def main():
    '''
        main routine
    '''
    print('[+] Starting project file running ')
    totallist = []
    commands = {"verbosemode": False,
                "deletemode": False,
                "startdirectory": TARGETDIRECTORY}
    totallist.extend(walk_through(commands))
    doclist = []
    print(f"[-] Records found {len(totallist):,}")

    for item in tqdm(totallist, total=len(totallist), unit=' item'):
        tmpext = extract_file_extension(item)
        if tmpext == '.docx':
            doclist.append(item)
    print(f'[-] Total DOC files is {len(doclist):,}')
    print('[-] Checking DOC files ')
    print(f'[-] Target string is {TARGETSTRING}')
    targetlist = []
    errorlist = []
    for docitem in tqdm(doclist, total=len(doclist), unit=' item'):
        try:
            doc_reader = docx.Document(docitem)
        except Exception as err:
            errorlist.append(f'[-] Error {err} in {docitem}')
            continue
        pagecount = len(doc_reader.paragraphs)
        for page in range(pagecount):
            text = doc_reader.paragraphs[page].text
            if TARGETSTRING in text:
                targetlist.append(docitem)
    print(f'[-] Total DOC files with target string is {len(targetlist):,}')
    if len(targetlist) > 0:
        for item in targetlist:
            print(f'[-] Found DOC file with target string {item}')
    if len(errorlist) > 0:
        for erritem in errorlist:
            print(f'[-] Error {erritem}')
    print('[+] Finishing project file running ')


if __name__ == '__main__':
    main()
