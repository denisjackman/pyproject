'''
    more word python fun
'''
import docx

# Reading a word doc
# create document object
doc = docx.Document("y:/Resources/word/my_word_file.docx")

all_paras = doc.paragraphs
print(len(all_paras))
for para in all_paras:
    print(para.text)
    print('-----------')


single_para = doc.paragraphs[2]
print(single_para.text)

for run in single_para.runs:
    print(run.text)

# Writing a word doc
mydoc = docx.Document()
mydoc.add_paragraph("This is first paragraph of a MS Word file.")
mydoc.save("y:/Resources/word/newmy_written_file_new.docx")
mydoc.add_paragraph("This is the second paragraph of a MS Word file.")
mydoc.save("y:/Resources/word/newmy_written_file_new.docx")
third_para = mydoc.add_paragraph("This is the third paragraph.")
third_para.add_run(" this is a section at the end of third paragraph")
mydoc.save("y:/Resources/word/newmy_written_file_new.docx")
mydoc.add_heading("This is level 1 heading", 0)
mydoc.add_heading("This is level 2 heading", 1)
mydoc.add_heading("This is level 3 heading", 2)
mydoc.save("y:/Resources/word/newmy_written_file_new.docx")
mydoc.add_picture("Z:/Resources/images/Casper.png",
                  width=docx.shared.Inches(5),
                  height=docx.shared.Inches(7))
mydoc.save("y:/Resources/word/newmy_written_file_new.docx")
